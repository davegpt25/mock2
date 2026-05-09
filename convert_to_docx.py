from docx import Document
from docx.shared import Pt, RGBColor, Inches, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_ALIGN_VERTICAL
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import copy

doc = Document()

# ── 기본 스타일 설정 ──────────────────────────────────────────────
style = doc.styles['Normal']
style.font.name = '맑은 고딕'
style.font.size = Pt(10.5)

def set_font(run, bold=False, size=None, color=None, name='맑은 고딕'):
    run.font.name = name
    run._element.rPr.rFonts.set(qn('w:eastAsia'), name)
    if bold:   run.font.bold = True
    if size:   run.font.size = Pt(size)
    if color:  run.font.color.rgb = RGBColor(*color)

def heading(doc, text, level, color=(0,0,0)):
    p = doc.add_paragraph()
    p.style = f'Heading {level}'
    run = p.add_run(text)
    run.font.name = '맑은 고딕'
    run._element.rPr.rFonts.set(qn('w:eastAsia'), '맑은 고딕')
    run.font.color.rgb = RGBColor(*color)
    return p

def bullet(doc, text, bold_part=None):
    p = doc.add_paragraph(style='List Bullet')
    p.paragraph_format.left_indent = Inches(0.25)
    if bold_part and bold_part in text:
        idx = text.index(bold_part)
        end = idx + len(bold_part)
        if idx > 0:
            r = p.add_run(text[:idx])
            set_font(r)
        rb = p.add_run(bold_part)
        set_font(rb, bold=True)
        rest = text[end:]
        if rest:
            r = p.add_run(rest)
            set_font(r)
    else:
        r = p.add_run(text)
        set_font(r)
    return p

def add_table(doc, headers, rows, col_widths=None):
    table = doc.add_table(rows=1+len(rows), cols=len(headers))
    table.style = 'Table Grid'
    # 헤더
    hdr = table.rows[0]
    hdr.height = Cm(0.8)
    for i, h in enumerate(headers):
        cell = hdr.cells[i]
        cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
        shd = OxmlElement('w:shd')
        shd.set(qn('w:val'), 'clear')
        shd.set(qn('w:color'), 'auto')
        shd.set(qn('w:fill'), '2E75B6')
        cell._tc.get_or_add_tcPr().append(shd)
        p = cell.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run(h)
        run.font.bold = True
        run.font.color.rgb = RGBColor(255, 255, 255)
        run.font.name = '맑은 고딕'
        run._element.rPr.rFonts.set(qn('w:eastAsia'), '맑은 고딕')
    # 데이터
    for ri, row in enumerate(rows):
        fill = 'F2F7FC' if ri % 2 == 0 else 'FFFFFF'
        for ci, val in enumerate(row):
            cell = table.rows[ri+1].cells[ci]
            cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
            shd = OxmlElement('w:shd')
            shd.set(qn('w:val'), 'clear')
            shd.set(qn('w:color'), 'auto')
            shd.set(qn('w:fill'), fill)
            cell._tc.get_or_add_tcPr().append(shd)
            p = cell.paragraphs[0]
            run = p.add_run(val)
            run.font.name = '맑은 고딕'
            run._element.rPr.rFonts.set(qn('w:eastAsia'), '맑은 고딕')
    if col_widths:
        for i, row in enumerate(table.rows):
            for j, cell in enumerate(row.cells):
                cell.width = Cm(col_widths[j])
    return table

def add_space(doc, size=6):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(size)
    p.paragraph_format.space_after  = Pt(0)

def hr(doc):
    p = doc.add_paragraph()
    pPr = p._p.get_or_add_pPr()
    pBdr = OxmlElement('w:pBdr')
    bottom = OxmlElement('w:bottom')
    bottom.set(qn('w:val'), 'single')
    bottom.set(qn('w:sz'), '6')
    bottom.set(qn('w:space'), '1')
    bottom.set(qn('w:color'), '2E75B6')
    pBdr.append(bottom)
    pPr.append(pBdr)

# ═══════════════════════════════════════════════════════════════════
# 표지
# ═══════════════════════════════════════════════════════════════════
doc.add_paragraph()
doc.add_paragraph()
p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = p.add_run('클로젯핏 (ClosetFit)')
r.font.name = '맑은 고딕'; r._element.rPr.rFonts.set(qn('w:eastAsia'), '맑은 고딕')
r.font.size = Pt(28); r.font.bold = True
r.font.color.rgb = RGBColor(0x2E, 0x75, 0xB6)

p2 = doc.add_paragraph()
p2.alignment = WD_ALIGN_PARAGRAPH.CENTER
r2 = p2.add_run('서비스 기획서')
r2.font.name = '맑은 고딕'; r2._element.rPr.rFonts.set(qn('w:eastAsia'), '맑은 고딕')
r2.font.size = Pt(18); r2.font.color.rgb = RGBColor(0x40, 0x40, 0x40)

doc.add_paragraph()
p3 = doc.add_paragraph()
p3.alignment = WD_ALIGN_PARAGRAPH.CENTER
r3 = p3.add_run('"이미 가진 옷으로, 더 잘 입는다"')
r3.font.name = '맑은 고딕'; r3._element.rPr.rFonts.set(qn('w:eastAsia'), '맑은 고딕')
r3.font.size = Pt(13); r3.font.italic = True
r3.font.color.rgb = RGBColor(0x70, 0x70, 0x70)

doc.add_paragraph()
p4 = doc.add_paragraph()
p4.alignment = WD_ALIGN_PARAGRAPH.CENTER
r4 = p4.add_run('2026. 05. 09')
r4.font.name = '맑은 고딕'; r4._element.rPr.rFonts.set(qn('w:eastAsia'), '맑은 고딕')
r4.font.size = Pt(11); r4.font.color.rgb = RGBColor(0x90, 0x90, 0x90)

doc.add_page_break()

# ═══════════════════════════════════════════════════════════════════
# 1. 서비스 개요
# ═══════════════════════════════════════════════════════════════════
heading(doc, '1. 서비스 개요', 1, (0x2E, 0x75, 0xB6))
hr(doc)
add_space(doc)

add_table(doc,
    ['항목', '내용'],
    [
        ['서비스명',      '클로젯핏 (ClosetFit)'],
        ['서비스 유형',   '모바일 앱 (iOS / Android)'],
        ['핵심 가치',     '"이미 가진 옷으로, 더 잘 입는다"'],
        ['타겟 사용자',   '20~40대 패션에 관심 있지만 코디가 어려운 남녀'],
    ],
    col_widths=[4, 12]
)

add_space(doc)
heading(doc, '서비스 한 줄 정의', 2, (0x2E, 0x75, 0xB6))
p = doc.add_paragraph()
p.paragraph_format.left_indent = Inches(0.3)
r = p.add_run('내 옷장에 있는 상의·하의를 선택하면, AI가 어울리는 컬러 팔레트와 추천 룩을 제시하여 스타일링을 도와주는 퍼스널 코디 서비스')
set_font(r, bold=True)

# ═══════════════════════════════════════════════════════════════════
# 2. 문제 정의
# ═══════════════════════════════════════════════════════════════════
add_space(doc, 12)
heading(doc, '2. 문제 정의 (Problem Statement)', 1, (0x2E, 0x75, 0xB6))
hr(doc)
add_space(doc)
heading(doc, '사용자 페인포인트', 2, (0x2E, 0x75, 0xB6))

pain_points = [
    ('"옷은 많은데 입을 게 없다"', ' — 옷장에 옷이 있어도 조합을 몰라 같은 옷만 입게 됨'),
    ('매칭 기준 부재',              ' — 컬러, 톤, 핏 등 코디 기준을 직관적으로 알기 어려움'),
    ('쇼핑 의존성',                 ' — 기존 옷을 활용하지 못하고 불필요한 소비로 이어짐'),
    ('레퍼런스 탐색 피로',          ' — 핀터레스트, 인스타그램 등에서 참고 룩을 찾는 데 시간 소요'),
]
for bold_part, rest in pain_points:
    p = doc.add_paragraph(style='List Bullet')
    p.paragraph_format.left_indent = Inches(0.25)
    rb = p.add_run(bold_part); set_font(rb, bold=True)
    rr = p.add_run(rest);       set_font(rr)

# ═══════════════════════════════════════════════════════════════════
# 3. 서비스 목표
# ═══════════════════════════════════════════════════════════════════
add_space(doc, 12)
heading(doc, '3. 서비스 목표', 1, (0x2E, 0x75, 0xB6))
hr(doc)
add_space(doc)
goals = [
    '보유 의류 기반의 코디 매칭 정확도 향상',
    '컬러 중심의 직관적인 스타일 가이드 제공',
    '불필요한 쇼핑 충동 감소 → 지속 가능한 패션 소비 유도',
    '사용자별 퍼스널 컬러 및 선호 스타일 학습을 통한 개인화 강화',
]
for i, g in enumerate(goals, 1):
    p = doc.add_paragraph(style='List Number')
    p.paragraph_format.left_indent = Inches(0.25)
    r = p.add_run(g); set_font(r)

# ═══════════════════════════════════════════════════════════════════
# 4. 핵심 기능
# ═══════════════════════════════════════════════════════════════════
add_space(doc, 12)
heading(doc, '4. 핵심 기능 (Core Features)', 1, (0x2E, 0x75, 0xB6))
hr(doc)

features = [
    ('4-1. 내 옷 등록 (My Closet)', [
        '상의 / 하의 / 아우터 / 신발 / 액세서리 카테고리 구분',
        '사진 촬영 또는 갤러리 업로드',
        'AI 자동 태깅: 컬러, 소재 유형, 핏, 스타일 키워드 추출',
        '직접 수정 가능한 태그 편집 UI',
    ]),
    ('4-2. 컬러 매칭 추천', [
        '선택한 아이템의 주 컬러 / 보조 컬러 자동 추출',
        '컬러 조합 원칙 기반 추천: 유사 색상(Analogous) / 보색(Complementary) / 중립(Neutral) / 톤온톤',
        '추천 컬러 팔레트를 시각적으로 제시',
    ]),
    ('4-3. 룩 추천 (Outfit Suggestion)', [
        '등록된 내 옷 중에서 어울리는 조합을 2~5가지 제안',
        '스타일 무드 필터: 캐주얼 / 오피스 / 데이트 / 스트릿 / 미니멀',
        '계절·날씨 연동 필터',
        '추천 룩에 부족한 아이템 제안 (구매 연계)',
    ]),
    ('4-4. 레퍼런스 룩 갤러리', [
        '선택 아이템의 컬러·스타일 기반 실제 코디 사진 큐레이션',
        '스타일리스트·인플루언서 룩 태그 연계',
        '저장 / 공유 기능',
    ]),
    ('4-5. 퍼스널 스타일 리포트', [
        '자주 입는 컬러 분포, 선호 스타일 패턴 분석',
        '옷장 활용률 리포트 (잘 안 입는 옷 알림)',
        '퍼스널 컬러 진단 연동 (봄웜·여름쿨·가을웜·겨울쿨)',
    ]),
]
for title, items in features:
    add_space(doc)
    heading(doc, title, 2, (0x2E, 0x75, 0xB6))
    for item in items:
        p = doc.add_paragraph(style='List Bullet')
        p.paragraph_format.left_indent = Inches(0.25)
        r = p.add_run(item); set_font(r)

# ═══════════════════════════════════════════════════════════════════
# 5. 사용자 플로우
# ═══════════════════════════════════════════════════════════════════
add_space(doc, 12)
heading(doc, '5. 사용자 플로우 (User Flow)', 1, (0x2E, 0x75, 0xB6))
hr(doc)
add_space(doc)

flow_lines = [
    '앱 실행',
    '  └─ 온보딩: 퍼스널 컬러 / 선호 스타일 설정',
    '       └─ 홈 화면',
    '            ├─ [오늘 뭐 입지?] → 상의 또는 하의 선택',
    '            │      └─ 컬러 팔레트 분석 → 매칭 아이템 추천 → 룩 완성',
    '            │              └─ 레퍼런스 룩 보기 → 저장 / 공유',
    '            ├─ [내 옷장] → 아이템 등록 / 관리',
    '            └─ [스타일 리포트] → 컬러 분석 / 활용률 확인',
]
for line in flow_lines:
    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Inches(0.2)
    r = p.add_run(line)
    r.font.name = 'Courier New'
    r.font.size = Pt(9.5)

# ═══════════════════════════════════════════════════════════════════
# 6. 화면 구성
# ═══════════════════════════════════════════════════════════════════
add_space(doc, 12)
heading(doc, '6. 화면 구성 (주요 화면)', 1, (0x2E, 0x75, 0xB6))
hr(doc)
add_space(doc)
add_table(doc,
    ['화면', '설명'],
    [
        ['온보딩',        '퍼스널 컬러 진단, 선호 스타일 선택 (최대 3가지)'],
        ['홈',            '오늘의 날씨 기반 코디 제안, 최근 착용 룩'],
        ['코디 선택',     '상의/하의 아이템 선택 → 매칭 분석 진입'],
        ['컬러 매칭 결과','추출된 컬러 팔레트 + 추천 컬러 조합 시각화'],
        ['룩 추천',       '내 옷장 기반 조합 카드 (스와이프) + 레퍼런스 룩'],
        ['내 옷장',       '카테고리별 아이템 그리드, 등록/삭제/태그 수정'],
        ['스타일 리포트', '주간/월간 컬러 통계, 미착용 아이템 알림'],
    ],
    col_widths=[4, 12]
)

# ═══════════════════════════════════════════════════════════════════
# 7. 기술 스택
# ═══════════════════════════════════════════════════════════════════
add_space(doc, 12)
heading(doc, '7. 기술 스택 (안)', 1, (0x2E, 0x75, 0xB6))
hr(doc)
add_space(doc)
add_table(doc,
    ['영역', '기술'],
    [
        ['프론트엔드',  'React Native (iOS / Android 동시 지원)'],
        ['백엔드',      'Node.js + FastAPI (AI 서빙 분리)'],
        ['AI / ML',     '컬러 추출: OpenCV + K-means 클러스터링'],
        ['',            '이미지 태깅: CLIP / Vision API'],
        ['',            '추천 엔진: 컬러 이론 룰베이스 + 협업 필터링'],
        ['데이터베이스','PostgreSQL (사용자/아이템), Redis (세션/캐싱)'],
        ['스토리지',    'AWS S3 (의류 이미지)'],
        ['인증',        '소셜 로그인 (카카오, 애플, 구글)'],
    ],
    col_widths=[4, 12]
)

# ═══════════════════════════════════════════════════════════════════
# 8. 수익 모델
# ═══════════════════════════════════════════════════════════════════
add_space(doc, 12)
heading(doc, '8. 수익 모델 (Monetization)', 1, (0x2E, 0x75, 0xB6))
hr(doc)
add_space(doc)
add_table(doc,
    ['모델', '내용'],
    [
        ['프리미엄 구독', '무제한 아이템 등록, 고급 스타일 리포트, 광고 제거'],
        ['제휴 커머스',   '추천 룩 내 부족 아이템 연계 쇼핑 (CPS 수수료)'],
        ['브랜드 콜라보', '시즌별 스타일 가이드 스폰서십'],
        ['데이터 B2B',    '익명화된 컬러 트렌드 데이터 패션 브랜드 제공'],
    ],
    col_widths=[4, 12]
)

# ═══════════════════════════════════════════════════════════════════
# 9. 경쟁 서비스 분석
# ═══════════════════════════════════════════════════════════════════
add_space(doc, 12)
heading(doc, '9. 경쟁 서비스 분석', 1, (0x2E, 0x75, 0xB6))
hr(doc)
add_space(doc)
add_table(doc,
    ['서비스', '특징', '차별점 (vs 클로젯핏)'],
    [
        ['클리오 (Clio)',     '옷장 관리 앱',           '코디 추천 기능 약함'],
        ['스타일쉐어',        '커뮤니티 기반 스타일',   '내 옷 기반 매칭 없음'],
        ['21버튼',            '쇼핑 연계 룩북',         '신규 구매 중심'],
        ['ChatGPT 패션 활용', '텍스트 기반 조언',       '이미지 기반 직관적 UX 부재'],
    ],
    col_widths=[3.5, 6, 6.5]
)
add_space(doc)
p = doc.add_paragraph()
r = p.add_run('클로젯핏의 핵심 차별화: ')
set_font(r, bold=True)
r2 = p.add_run('내가 이미 가진 옷을 기반으로 컬러 과학을 적용한 실용적 코디 가이드')
set_font(r2)

# ═══════════════════════════════════════════════════════════════════
# 10. 출시 로드맵
# ═══════════════════════════════════════════════════════════════════
add_space(doc, 12)
heading(doc, '10. 출시 로드맵', 1, (0x2E, 0x75, 0xB6))
hr(doc)
add_space(doc)
add_table(doc,
    ['단계', '기간', '주요 내용'],
    [
        ['Phase 1 — MVP',        '1~3개월',   '옷 등록, 컬러 추출, 기본 매칭 추천, 레퍼런스 룩'],
        ['Phase 2 — 개인화',     '4~6개월',   '퍼스널 컬러 진단, 스타일 필터, 리포트 기능'],
        ['Phase 3 — 커머스 연동','7~9개월',   '부족 아이템 쇼핑 연계, 브랜드 파트너십'],
        ['Phase 4 — 커뮤니티',   '10~12개월', '유저 룩 공유, 팔로우, 코디 피드'],
    ],
    col_widths=[4.5, 2.5, 9]
)

# ═══════════════════════════════════════════════════════════════════
# 11. KPI
# ═══════════════════════════════════════════════════════════════════
add_space(doc, 12)
heading(doc, '11. KPI (핵심 성과 지표)', 1, (0x2E, 0x75, 0xB6))
hr(doc)
add_space(doc)
kpis = [
    ('DAU/MAU 비율',    '목표: 30% 이상 (매일 사용하는 습관형 앱 지향)'),
    ('아이템 등록 수',  '사용자당 평균 20개 이상'),
    ('코디 추천 사용률','주 3회 이상 사용 비율'),
    ('구독 전환율',     '프리미엄 전환 5% 이상 (6개월 내)'),
    ('커머스 클릭률',   '추천 아이템 클릭 10% 이상'),
]
for bold_part, rest in kpis:
    p = doc.add_paragraph(style='List Bullet')
    p.paragraph_format.left_indent = Inches(0.25)
    rb = p.add_run(bold_part + ' — '); set_font(rb, bold=True)
    rr = p.add_run(rest);              set_font(rr)

# ═══════════════════════════════════════════════════════════════════
# 12. 리스크 및 대응
# ═══════════════════════════════════════════════════════════════════
add_space(doc, 12)
heading(doc, '12. 리스크 및 대응', 1, (0x2E, 0x75, 0xB6))
hr(doc)
add_space(doc)
add_table(doc,
    ['리스크', '대응 방안'],
    [
        ['옷 등록 번거로움으로 인한 이탈',    '온보딩 시 10개 이하 최소 등록 유도, 배경 제거 자동화'],
        ['컬러 추출 정확도 저하 (조명 차이)', '촬영 가이드 UI 제공, 사용자 수동 보정 옵션'],
        ['추천 다양성 부족',                  '스타일 무드별 필터 + 주기적 레퍼런스 데이터 업데이트'],
        ['개인정보 (의류 이미지) 우려',       '이미지 서버 암호화, 삭제 요청 즉시 처리 정책'],
    ],
    col_widths=[6, 10]
)

# ── 저장 ──────────────────────────────────────────────────────────
output = r'C:\Users\hwlll\Downloads\기획서\클로젯핏_서비스기획서.docx'
doc.save(output)
print(f'저장 완료: {output}')
